from docx import Document
from docx.shared import Inches, Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime as dt



class docx_generator():

    def __init__(self):
        self.main_font_size = Pt(12.0)
        self.doc = Document()

    def add_page_measurements(self):
        # set page height and width and ...
        section = self.doc.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Mm(5)
        section.bottom_margin = Mm(5)

    def add_heading(self):
        head_text = 'ROYAUME DU MAROC\nMINISTERE DE L\'INTERIEUR\nWILAYA DE LA REGION DE RABAT-\nSALE-KENITRA\nPREFECTURE DE RABAT\nSECRETARIAT GENERAL\nDIVISION DU BUDGET ET DES MARCHES\nSERVICE DU BUDGET'
        heading = self.doc.add_heading(head_text,3)
        heading.paragraph_format.left_indent = Inches(-5.0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_title(self):
        title = self.doc.add_paragraph()
        title_run = title.add_run('ATTESTATION DE DOMICILIATION\nDE SALAIRE IRREVOCABLE\n*****************')
        title_run.font.size = Pt(16.0)
        title_run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Inches(0.5)

    def add_name(self, n ='walo', s = 'm'):
        if s == 'm':
            tmp = 'M'
        elif s == 'f':
            tmp = 'Mme'
        name = self.doc.add_paragraph()
        name.add_run(f'Le Wali de la Région de Rabat-Salé-Kenitra, Gouverneur de la Préfecture de Rabat, sous-ordonnateur du Budget Général. Atteste par la présente que le salaire mensuel de {tmp}:  ').font.size = self.main_font_size
        name.add_run(n).font.size = self.main_font_size
        name.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        name.paragraph_format.first_line_indent = Inches(0.5)

    def add_grade(self, g = 'walo'):
        grade = self.doc.add_paragraph()
        grade.add_run('Grade:  ').font.size = self.main_font_size
        grade.add_run(g).font.size = self.main_font_size

    def add_workplace(self, w = 'walo'):
        workplace = self.doc.add_paragraph()
        workplace.add_run('En fonction au :  ').font.size = self.main_font_size
        if w == 'Autre': w = '..............................'
        workplace.add_run(w).font.size = self.main_font_size


    def rib_text(func):

        def inner_func(self, *args):
            rib_text = self.doc.add_paragraph()
            rib_text.add_run('Sera à compter de ce jour viré chaque mois à son compte  Bancaire N° :  ').font.size = self.main_font_size
            rib_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            func(self, *args)
        
        return inner_func
    
    @rib_text
    def add_rib(self, r = 'walo'):
        rib = self.doc.add_paragraph()
        rib.add_run(r).font.size = self.main_font_size
        rib.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_bank(self, b = 'walo'):
        bank = self.doc.add_paragraph()
        bank.add_run('Ouvert auprès de :  ').font.size = self.main_font_size
        bank.add_run(b).font.size = self.main_font_size

    def add_agence(self, a = 'walo', s = 'm'):
        if s == 'm':
            tmp = ''
        elif s == 'f':
            tmp = 'e'

        agence = self.doc.add_paragraph()
        agence.add_run('Agence de :  ').font.size = self.main_font_size
        agence.add_run('.' * 30).font.size = self.main_font_size
        agence.add_run(f'  Tant qu’il est en exercice dans sa fonction.\nPar ailleurs, aucune  avance  sur salaire  ne sera versée à l\'intéressé{tmp} au cas où  celui-ci cesserait  son activité au sein de cette préfecture, cette dernière décline sa responsabilité en cas de démission, de révocation ou du décès de l’intéressé.').font.size = self.main_font_size
        agence.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def end_of_file(self):
        note = self.doc.add_paragraph()
        note.add_run('Le solde des émoluments sera versé au  compte bancaire susmentionné.').font.size = self.main_font_size
        note.paragraph_format.first_line_indent = Inches(0.5)
        note.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
        note.paragraph_format.line_spacing = 2.0

        date = self.doc.add_paragraph()
        date.add_run('Rabat, le: ').font.size = Pt(13.0)
        date_run = date.add_run(dt.datetime.strftime(dt.datetime.now(),'%d %b %Y'))
        date_run.font.size = Pt(13.0)
        date_run.font.color.rgb = RGBColor(0,0,255)
        date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date.paragraph_format.line_spacing = 2.0


        signature = self.doc.add_table(rows=1, cols=2)
        signature.alignment = WD_TABLE_ALIGNMENT.CENTER
        cells = signature.rows[0].cells
        cells[0].text = 'Accord du salaire et signature  '
        cells[1].text = 'LE SOUS-ORDONNATEUR'


        # change the text size for the cells[0] & cells[1]
        i = 0
        for cell in cells:
            if i == 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                i += 1
            else:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = cell.paragraphs[0].runs
            run[0].font.size = self.main_font_size


